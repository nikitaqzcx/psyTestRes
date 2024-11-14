from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from typing import List


uk_months = ["Січень", "Лютий", "Березень",
     "Квітень",  "Травень",  "Червень",
     "Липень",  "Серпень", "Вересень",
     "Жовтень","Листопад","Грудень"]


class SingleResult:
    def __init__(self, name: str, date: datetime, info: List[str]):
        self.name = name
        self.date = date
        self.info = info


class LasarusResults:
    def __init__(self, file_name: str, date_from: datetime = None, date_to: datetime = None):
        self.file_name = file_name
        self.error_rows = []
        self.use_date_filter = False
        self.date_from = date_from
        self.date_to = date_to

        # Enable date filtering if a valid date range is provided
        if self.date_from and self.date_to:
            self.use_date_filter = True


    def create_date_break(self, doc, text):
        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles['Heading 1']
        run = paragraph.add_run(text)
        run.font.size = Pt(16)
        run.font.name = "Calibri Light"
        run.font.color.rgb = RGBColor(47, 84, 150)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    def create_new_block(self, doc, text):
        doc.add_paragraph().add_run()
        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles['Heading 2']
        run = paragraph.add_run(text)
        run.font.size = Pt(16)
        run.font.name = "Calibri Light"
        run.font.color.rgb = RGBColor(47, 84, 150)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(0)

    def create_paragraph(self, doc, text, bold_first_sentence):
        paragraph = doc.add_paragraph()
        first_sentence_end = text.find('.') + 1
        if bold_first_sentence and first_sentence_end > 0:
            paragraph.add_run(text[:first_sentence_end]).bold = True
            paragraph.add_run(text[first_sentence_end:])
        else:
            paragraph.add_run(text)
        paragraph.paragraph_format.first_line_indent = Pt(12)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def add_child_block(self, doc, text, bold_first_sentence=True):
        paragraphs = text.split("|")
        for i, paragraph in enumerate(paragraphs):
            self.create_paragraph(doc, paragraph, bold_first_sentence if i == 0 else False)

    def cell_by_name(self, name):
        result = 0
        for char in name.strip().upper():
            result = result * 26 + (ord(char) - ord('A') + 1)
        return result - 1

    def gather_by_structure(self, sheet_name, structure):
        wb = openpyxl.load_workbook(self.file_name)
        sheet = wb[sheet_name]
        results = []
        for row_idx in range(2, sheet.max_row + 1):
            try:
                date_cell = sheet.cell(row=row_idx, column=self.cell_by_name(structure['date_column']) + 1).value
                if not isinstance(date_cell, datetime) or date_cell < datetime(1900, 1, 1):
                    raise ValueError(f"Invalid or missing date at row {row_idx}")

                # Only apply date filtering if the date range is provided
                if self.use_date_filter:
                    if self.date_from and date_cell < self.date_from:
                        continue
                    if self.date_to and date_cell > self.date_to:
                        continue

                name = sheet.cell(row=row_idx, column=self.cell_by_name(structure['name_column']) + 1).value
                if not name:
                    raise ValueError(f"Missing name at row {row_idx}")

                info = []
                for col in structure['data_columns'].split(','):
                    cell_value = sheet.cell(row=row_idx, column=self.cell_by_name(col) + 1).value
                    if cell_value is None:
                        raise ValueError(f"Missing data in column {col} at row {row_idx}")
                    info.append(cell_value)
                results.append(SingleResult(name=name, date=date_cell, info=info))
            
            except Exception as e:
                print(f"Error at row {row_idx}: {e}")
                self.error_rows.append(row_idx)
        
        return results

    @staticmethod
    def add_toc(doc):
        """Adds a Table of Contents (TOC) to the document."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Specify heading levels and options

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."  # Placeholder text for the TOC

        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)

    def compose_doc_by_structure(self, sheet_name, structure):
        doc = Document()
        results = sorted(self.gather_by_structure(sheet_name, structure), key=lambda x: (x.date, x.name))
        self.add_toc(doc)

        if results:
            month_index = int(results[0].date.strftime('%m')) - 1 
            curr_month = f"{uk_months[month_index]} {results[0].date.strftime('%Y')}"

            self.create_date_break(doc, curr_month)
            doc.add_page_break()
            
            for result in results:
                next_month_index = int(result.date.strftime('%m')) - 1 
                next_month = f"{uk_months[next_month_index]} {results[0].date.strftime('%Y')}"
                if next_month != curr_month:
                    self.create_date_break(doc, next_month)
                    curr_month = next_month

                self.create_new_block(doc, result.name)
                # Format date as DD-MM-YYYY
                self.add_child_block(doc, result.date.strftime('%d-%m-%Y %H:%M:%S'), False)  # Date in DD-MM-YYYY format
                for part in result.info:
                    self.add_child_block(doc, part)
               
                doc.add_page_break()
                
        return doc

    def save_results(self, sheet_name, output_file_name, structure):
        try:
            doc = self.compose_doc_by_structure(sheet_name, structure)
            doc.save(output_file_name)
            if not self.error_rows:
                print("Saved successfully!")
            else:
                print(f"Document saved, but the following rows could not be processed: {', '.join(map(str, self.error_rows))}")
        except Exception as e:
            print("Failed to save the document:", e)

